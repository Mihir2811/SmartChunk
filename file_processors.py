# Handles different file formats and extraction
from abc import ABC, abstractmethod
from typing import List, Dict, Any
import json
import yaml
from pathlib import Path

from pypdf import PdfReader
from docx import Document

from utils import get_logger

logger = get_logger(__name__)

class FileProcessor(ABC):
    """Abstract base class for file processors."""
    
    @abstractmethod
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the file type."""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract text content from the file."""
        pass
    
    @abstractmethod
    def extract_structured_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract structured sections with metadata.
        
        Returns:
            List of dicts with 'content', 'section_name', and other metadata
        """
        pass

class PDFProcessor(FileProcessor):
    """Processor for PDF files."""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')
    
    def extract_text(self, file_path: str) -> str:
        logger.info(f"Extracting text from PDF: {file_path}")
        
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{text}")
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from {len(reader.pages)} pages")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    def extract_structured_sections(self, file_path: str) -> List[Dict[str, Any]]:
        reader = PdfReader(file_path)
        sections = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                sections.append({
                    "content": text,
                    "section_name": f"Page {page_num}",
                    "page_number": page_num
                })
        
        return sections

class WordProcessor(FileProcessor):
    """Processor for Word documents (.docx, .doc)."""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def extract_text(self, file_path: str) -> str:
        logger.info(f"Extracting text from Word document: {file_path}")
        
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            full_text = "\n\n".join(paragraphs)
            logger.info(f"Extracted {len(full_text)} characters from {len(paragraphs)} paragraphs")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting Word document: {e}")
            raise
    
    def extract_structured_sections(self, file_path: str) -> List[Dict[str, Any]]:
        doc = Document(file_path)
        sections = []
        current_section = None
        current_content = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # Detect headings based on style
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_section and current_content:
                    sections.append({
                        "content": "\n\n".join(current_content),
                        "section_name": current_section
                    })
                
                # Start new section
                current_section = para.text
                current_content = []
            else:
                current_content.append(para.text)
        
        # Add final section
        if current_section and current_content:
            sections.append({
                "content": "\n\n".join(current_content),
                "section_name": current_section
            })
        
        return sections if sections else [{"content": self.extract_text(file_path), "section_name": "Document"}]

class MarkdownProcessor(FileProcessor):
    """Processor for Markdown files."""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.md', '.markdown'))
    
    def extract_text(self, file_path: str) -> str:
        logger.info(f"Extracting text from Markdown: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        logger.info(f"Extracted {len(content)} characters")
        return content
    
    def extract_structured_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract sections based on Markdown headers."""
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        sections = []
        current_header = None
        current_level = 0
        current_content = []
        header_stack = []
        
        for line in lines:
            # Check for headers
            if line.startswith('#'):
                # Save previous section
                if current_header and current_content:
                    sections.append({
                        "content": "".join(current_content).strip(),
                        "section_name": " > ".join(header_stack),
                        "header_level": current_level
                    })
                
                # Parse new header
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                
                # Update header stack
                while len(header_stack) >= level:
                    header_stack.pop()
                header_stack.append(header_text)
                
                current_header = header_text
                current_level = level
                current_content = []
            else:
                current_content.append(line)
        
        # Add final section
        if current_header and current_content:
            sections.append({
                "content": "".join(current_content).strip(),
                "section_name": " > ".join(header_stack),
                "header_level": current_level
            })
        
        return sections if sections else [{"content": self.extract_text(file_path), "section_name": "Document"}]

class JSONProcessor(FileProcessor):
    """Processor for JSON files, including Postman collections."""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.json')
    
    def extract_text(self, file_path: str) -> str:
        logger.info(f"Extracting text from JSON: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Pretty print for readability
        content = json.dumps(data, indent=2)
        logger.info(f"Extracted {len(content)} characters")
        
        return content
    
    def extract_structured_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract sections, with special handling for Postman collections."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        sections = []
        
        # Check if this is a Postman collection
        if self._is_postman_collection(data):
            sections = self._extract_postman_requests(data)
        else:
            # Generic JSON processing - split by top-level keys
            for key, value in data.items():
                sections.append({
                    "content": json.dumps({key: value}, indent=2),
                    "section_name": f"JSON Key: {key}",
                    "json_key": key
                })
        
        return sections if sections else [{"content": self.extract_text(file_path), "section_name": "JSON Document"}]
    
    def _is_postman_collection(self, data: dict) -> bool:
        """Check if JSON is a Postman collection."""
        return 'info' in data and 'item' in data and data.get('info', {}).get('schema', '').find('collection') != -1
    
    def _extract_postman_requests(self, data: dict) -> List[Dict[str, Any]]:
        """Extract individual requests from Postman collection."""
        sections = []
        collection_name = data.get('info', {}).get('name', 'Unknown Collection')
        
        def process_items(items, folder_path=""):
            for item in items:
                if 'item' in item:
                    # This is a folder
                    folder_name = item.get('name', 'Unnamed Folder')
                    new_path = f"{folder_path}/{folder_name}" if folder_path else folder_name
                    process_items(item['item'], new_path)
                else:
                    # This is a request
                    request_name = item.get('name', 'Unnamed Request')
                    request_data = item.get('request', {})
                    
                    content_parts = [
                        f"Request: {request_name}",
                        f"Method: {request_data.get('method', 'GET')}",
                        f"URL: {self._format_url(request_data.get('url', {}))}",
                    ]
                    
                    # Add description if available
                    if item.get('request', {}).get('description'):
                        content_parts.append(f"Description: {request_data['description']}")
                    
                    # Add headers
                    headers = request_data.get('header', [])
                    if headers:
                        content_parts.append("\nHeaders:")
                        for header in headers:
                            content_parts.append(f"  {header.get('key')}: {header.get('value')}")
                    
                    # Add body
                    body = request_data.get('body', {})
                    if body:
                        content_parts.append(f"\nBody ({body.get('mode', 'raw')}):")
                        content_parts.append(str(body.get(body.get('mode', 'raw'), '')))
                    
                    section_name = f"{collection_name}/{folder_path}/{request_name}" if folder_path else f"{collection_name}/{request_name}"
                    
                    sections.append({
                        "content": "\n".join(content_parts),
                        "section_name": section_name,
                        "request_method": request_data.get('method', 'GET'),
                        "request_name": request_name
                    })
        
        process_items(data.get('item', []))
        return sections
    
    def _format_url(self, url_data) -> str:
        """Format Postman URL object to string."""
        if isinstance(url_data, str):
            return url_data
        elif isinstance(url_data, dict):
            return url_data.get('raw', str(url_data))
        return str(url_data)

class YAMLProcessor(FileProcessor):
    """Processor for YAML files."""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.yaml', '.yml'))
    
    def extract_text(self, file_path: str) -> str:
        logger.info(f"Extracting text from YAML: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        logger.info(f"Extracted {len(content)} characters")
        return content
    
    def extract_structured_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract sections based on top-level YAML keys."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        sections = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                content = yaml.dump({key: value}, default_flow_style=False, sort_keys=False)
                sections.append({
                    "content": content,
                    "section_name": f"YAML Section: {key}",
                    "yaml_key": key
                })
        else:
            # Not a dictionary, treat as single section
            sections.append({
                "content": yaml.dump(data, default_flow_style=False),
                "section_name": "YAML Document"
            })
        
        return sections

class CodeProcessor(FileProcessor):
    """Processor for code files (Python, JavaScript, etc.)."""
    
    SUPPORTED_EXTENSIONS = {'.py', '.js', '.java', '.ts', '.tsx', '.jsx', '.html', '.css', '.cpp', '.c', '.go', '.rs'}
    
    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file_path: str) -> str:
        logger.info(f"Extracting code from: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        logger.info(f"Extracted {len(content)} characters")
        return content
    
    def extract_structured_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract code maintaining structure."""
        # For now, return whole file as one section
        # Future enhancement: parse AST and split by functions/classes
        return [{
            "content": self.extract_text(file_path),
            "section_name": Path(file_path).name,
            "language": Path(file_path).suffix.lstrip('.')
        }]

class PlainTextProcessor(FileProcessor):
    """Processor for plain text files."""
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.txt')
    
    def extract_text(self, file_path: str) -> str:
        logger.info(f"Extracting text from: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        logger.info(f"Extracted {len(content)} characters")
        return content
    
    def extract_structured_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """Split by double newlines (paragraphs)."""
        content = self.extract_text(file_path)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        sections = []
        for i, para in enumerate(paragraphs, 1):
            sections.append({
                "content": para,
                "section_name": f"Paragraph {i}"
            })
        
        return sections if sections else [{"content": content, "section_name": "Document"}]

class FileProcessorFactory:
    """Factory to get appropriate processor for a file."""
    
    def __init__(self):
        self.processors = [
            PDFProcessor(),
            WordProcessor(),
            MarkdownProcessor(),
            JSONProcessor(),
            YAMLProcessor(),
            CodeProcessor(),
            PlainTextProcessor(),
        ]
    
    def get_processor(self, file_path: str) -> FileProcessor:
        """
        Get appropriate processor for file type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            FileProcessor instance
            
        Raises:
            ValueError: If no processor found for file type
        """
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        
        raise ValueError(f"No processor found for file type: {Path(file_path).suffix}")